"""
Report Generator Utility
Converts research output (Markdown or dict) into a professional Word (.docx) file.
"""

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt

def generate_word_report(research, filename="Research_Report.docx"):
    doc = Document()
    doc.add_heading('Research Intelligence Report', 0)
    doc.add_paragraph('Prepared by: AI Research Agent (Ashish Singh Bora)')
    doc.add_page_break()
    doc.add_heading('Table of Contents', level=1)
    toc = doc.add_paragraph()
    for i, section in enumerate(research, 1):
        toc.add_run(f"{i}. {section['title']}\n").bold = True
    doc.add_page_break()
    for section in research:
        doc.add_heading(section['title'], level=1)
        doc.add_paragraph(section['content'])
    doc.add_page_break()
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('This report was generated automatically by the AI Agent. All sources are cited above.')
    doc.save(filename)
    return filename
