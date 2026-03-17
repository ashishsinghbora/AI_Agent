"""
Report Builder: Creates a professional .docx report from research summary and sources.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def build_word_report(summary: str, sources: list, filename: str = "Research_Report.docx"):
    doc = Document()
    # Title Header
    doc.add_heading('AI Research Report', 0)
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(summary)
    # Source Bibliography
    doc.add_heading('Source Bibliography', level=1)
    for src in sources:
        p = doc.add_paragraph()
        run = p.add_run(src['title'] + ': ')
        run.bold = True
        p.add_run(src['url'])
        # Make link clickable (Word will auto-link URLs)
    doc.save(filename)
    return filename
